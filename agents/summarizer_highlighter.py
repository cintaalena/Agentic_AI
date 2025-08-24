import os
import re
import spacy
import fitz
import docx
from docx.shared import RGBColor
import google.generativeai as genai 
from dotenv import load_dotenv


load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
else:
    print("PERINGATAN: GEMINI_API_KEY tidak ditemukan di file .env. Fitur ringkasan AI tidak akan berfungsi.")

print("Memuat model NLP (Spacy)...")
nlp_en = spacy.load("en_core_web_sm")
nlp_id = spacy.load("xx_ent_wiki_sm")
nlp_id.add_pipe('sentencizer')


project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
input_dir = os.path.join(project_root, 'input_files')
output_dir = os.path.join(project_root, 'output_files')
os.makedirs(output_dir, exist_ok=True)


# --- FUNGSI-FUNGSI UTAMA ---
def detect_language_from_text(text):
    """Mendeteksi bahasa (en/id) dari potongan teks di dalam dokumen."""
    from langdetect import detect
    try:
        snippet = text[:2000]
        lang = detect(snippet)
        return 'id' if lang == 'id' else 'en'
    except Exception:
        return 'en'

def extract_text_from_pdf(pdf_path):
    """Mengekstrak teks dari PDF menggunakan PyMuPDF (fitz) yang andal."""
    try:
        doc = fitz.open(pdf_path)
        text = "".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception as e:
        print(f"  -> Error saat mengekstrak PDF: {e}")
        return ""


def extract_text_from_docx(docx_path):
    """Mengekstrak teks dari DOCX."""
    try:
        doc = docx.Document(docx_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"   -> Gagal mengekstrak DOCX: {e}")
        return ""


def get_highly_relevant_sentences(text, language='en'):
    """
    (Versi 4 - Akurasi Tinggi) Logika yang sangat selektif untuk menyorot
    hanya bagian paling penting dari dokumen akademis.
    """
    print("  -> Menilai kalimat dengan logika akurasi tinggi (v4)...")
    nlp = nlp_id if language == 'id' else nlp_en
    
    SCORE_THRESHOLD = 3  
    MIN_SENTENCE_LENGTH_WORDS = 12 


    HIGH_IMPACT_KEYWORDS_ID = {'hasil', 'kesimpulan', 'metode', 'analisis', 'temuan', 'membuktikan', 'menunjukkan bahwa'}
    MEDIUM_IMPACT_KEYWORDS_ID = {'penelitian', 'tujuan', 'latar belakang', 'data', 'implikasi', 'hipotesis', 'evaluasi', 'responden', 'dampak', 'signifikan'}
    
    HIGH_IMPACT_KEYWORDS_EN = {'result', 'conclusion', 'method', 'analysis', 'finding', 'proves', 'shows that'}
    MEDIUM_IMPACT_KEYWORDS_EN = {'research', 'objective', 'background', 'data', 'implication', 'hypothesis', 'evaluation', 'respondent', 'impact', 'significant'}
    
    high_keywords = HIGH_IMPACT_KEYWORDS_ID if language == 'id' else HIGH_IMPACT_KEYWORDS_EN
    medium_keywords = MEDIUM_IMPACT_KEYWORDS_ID if language == 'id' else MEDIUM_IMPACT_KEYWORDS_EN

    IMPORTANT_ENTITY_TYPES = {"PERSON", "ORG", "PRODUCT", "EVENT", "LAW", "FAC", "LOC", "GPE"}
    DATA_ENTITY_TYPES = {"CARDINAL", "MONEY", "QUANTITY", "PERCENT"}
    
    relevant_sentences = set()

    # --- ATURAN 1 (DIPERBAIKI): Analisis Struktural dengan Filter ---
    lines = text.split('\n')
    HEADING_PATTERN = re.compile(
        r'^\s*(BAB\s+[IVXLCDM]+|LATAR BELAKANG|RUMUSAN PERMASALAHAN|TUJUAN DAN MANFAAT|KAJIAN KEPUSTAKAAN|METODOLOGI PENELITIAN|PENGEMBANGAN APLIKASI|PENGUJIAN|HASIL DAN PEMBAHASAN|KESIMPULAN)\s*$', 
        re.IGNORECASE
    )

    for i, line in enumerate(lines):
        if HEADING_PATTERN.match(line.strip()):
            for next_line in lines[i+1:]:
                if next_line.strip():
                    doc_line = nlp(next_line.strip())
                    first_sentence = next(doc_line.sents, None)
                    # Filter: Hanya ambil kalimat pertama jika cukup panjang & substantif
                    if first_sentence and len(first_sentence.text.strip().split()) > 8:
                        relevant_sentences.add(first_sentence.text.strip())
                    break 

    
    doc = nlp(text)
    all_sentences = list(doc.sents)
    
    
    start_index = 0
    BOILERPLATE_KEYWORDS = {'abstrak', 'kata kunci', 'daftar isi', 'lembar pengesahan', 'kata pengantar', 'ucapan terima kasih', 'npm', 'jurusan', 'program studi', 'tugas akhir'}
    for i, sent in enumerate(all_sentences[:30]): # Cek 30 kalimat pertama
        sent_lower = sent.text.lower()
        if len(sent.text.strip().split()) < 10 or any(kw in sent_lower for kw in BOILERPLATE_KEYWORDS):
            start_index = i + 1
        
    for sent in all_sentences[start_index:]:
        score = 0
        text_lower = sent.text.lower()
        
        
        if len(sent.text.split()) > MIN_SENTENCE_LENGTH_WORDS: score += 1
        if any(ent.label_ in IMPORTANT_ENTITY_TYPES for ent in sent.ents): score += 1
        if any(ent.label_ in DATA_ENTITY_TYPES for ent in sent.ents): score += 1

        
        if any(keyword in text_lower for keyword in high_keywords): score += 2
        elif any(keyword in text_lower for keyword in medium_keywords): score += 1

        if score >= SCORE_THRESHOLD:
            relevant_sentences.add(sent.text.strip())
            
    print(f"  -> Total {len(relevant_sentences)} kalimat AKURAT ditemukan untuk disorot.")
    return list(relevant_sentences)


def highlight_pdf_file(pdf_path, output_path, language='en'):
    """Membuat salinan PDF dengan sorotan visual pada kalimat yang paling relevan."""
    print(f"  -> Membuat PDF dengan sorotan untuk: {os.path.basename(pdf_path)}")
    doc = fitz.open(pdf_path)
    full_text = "".join(page.get_text() for page in doc)
    sentences_to_highlight = get_highly_relevant_sentences(full_text, language)

    if not sentences_to_highlight:
        print("  -> Tidak ada kalimat relevan untuk disorot. Menyimpan salinan asli.")
        doc.save(output_path)
        doc.close()
        return

    print(f"  -> Menyorot {len(sentences_to_highlight)} kalimat pada PDF...")
    for page in doc:
        for text_fragment in sentences_to_highlight:
            for inst in page.search_for(text_fragment):
                page.add_highlight_annot(inst).update()
                
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()
    print(f"  -> PDF dengan sorotan disimpan ke {os.path.basename(output_path)}")


def highlight_word_file(docx_path, output_path, language='en'):
    """
    Membuat salinan DOCX dengan penandaan visual pada paragraf relevan.
    CATATAN: Proses ini mungkin menghilangkan gambar dan format kompleks lainnya.
    """
    print(f"  -> Membuat DOCX dengan sorotan untuk: {os.path.basename(docx_path)}")
    try:
        original_doc = docx.Document(docx_path)
    except Exception as e:
        print(f"  -> Gagal membuka file DOCX: {e}")
        return

    full_text = "\n".join([p.text for p in original_doc.paragraphs])
    sentences_to_highlight = get_highly_relevant_sentences(full_text, language)

    if not sentences_to_highlight:
        print("  -> Tidak ada kalimat relevan untuk ditandai. Menyimpan salinan asli.")
        original_doc.save(output_path)
        return

    print(f"  -> Menyorot paragraf yang mengandung {len(sentences_to_highlight)} kalimat relevan...")
    
    for para in original_doc.paragraphs:
        if any(sent in para.text for sent in sentences_to_highlight):
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
    
    original_doc.save(output_path)
    print(f"  -> DOCX dengan sorotan disimpan ke {os.path.basename(output_path)}")
    

def summarize_with_ai_model(text, language='en'):
    """Membuat ringkasan "pintar" menggunakan Google Gemini API."""
    print("   -> Menganalisis teks untuk mendapatkan konteks sebelum meringkas...")
    
    if not GEMINI_API_KEY:
        return "Gagal: Kunci API Gemini tidak ditemukan di file .env."
    if len(text.strip().split()) < 150: 
        return "Gagal: Teks terlalu pendek untuk diringkas secara efektif oleh AI."
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        
        prompt = (
            "Anda adalah seorang analis riset ahli. Tugas Anda adalah membaca teks berikut dan membuat ringkasan analitis yang komprehensif.\n\n"
            "**Instruksi Spesifik:**\n"
            "1.  **Kedalaman**: Buat ringkasan yang mendalam dan informatif. JANGAN terlalu singkat. Tangkap poin-poin utama, metodologi, temuan kunci, dan kesimpulan.\n"
            "2.  **Struktur**: Awali dengan paragraf pembuka singkat, diikuti dengan poin-poin (bullet points) untuk detailnya.\n"
            "3.  **Bahasa**: Buat ringkasan dalam Bahasa "
            f"{'Indonesia' if language == 'id' else 'Inggris'}.\n\n"
            "--- TEKS DOKUMEN ---\n"
            f"{text[:25000]}"
            "\n\n--- AKHIR TEKS ---\n\n"
            "**Ringkasan Analitis Komprehensif:**"
        )
        
        response = model.generate_content(prompt)
        
        return response.text if response.parts else "Gagal: API Gemini tidak memberikan hasil."

    except Exception as e:
        print(f"   -> Terjadi error saat memanggil Gemini API: {e}")
        return f"Gagal: Terjadi error saat menghubungi layanan AI. ({e})"

def process_file(file_path):
    """Fungsi utama yang memproses satu file dan MENGEMBALIKAN path outputnya."""
    filename = os.path.basename(file_path)
    print(f"--- Memulai proses untuk: {filename} ---")
    text, file_type = "", ""
    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
        file_type = "pdf"
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_path)
        file_type = "docx"
    
    if not text or not text.strip(): 
        print(f"Tidak ada teks yang bisa diekstrak dari {filename}. Proses dihentikan."); 
        return None, None

    language = detect_language_from_text(text)
    print(f"   -> Bahasa terdeteksi: {language.upper()}")
    
    output_base_name = os.path.splitext(filename)[0]
    
    print("   -> Menggunakan model AI (Gemini API) untuk ringkasan pintar...")
    summary = summarize_with_ai_model(text, language)
    
    summary_path = os.path.join(output_dir, f"summary_{output_base_name}.txt")
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(summary)
    print(f"   -> Ringkasan AI disimpan ke {os.path.basename(summary_path)}")

    highlighted_path = None
    if file_type == "pdf":
        highlighted_path = os.path.join(output_dir, f"highlighted_{output_base_name}.pdf")
        highlight_pdf_file(file_path, highlighted_path, language)
    elif file_type == "docx":
        highlighted_path = os.path.join(output_dir, f"highlighted_{output_base_name}.docx")
        highlight_word_file(file_path, highlighted_path, language)
    
    print(f"--- Selesai memproses: {filename} ---\n")
    
    return summary_path, highlighted_path